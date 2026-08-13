"""Build the three Gemastik Word deliverables from repository-tracked sources."""

from __future__ import annotations

import re
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "gemastik-2026"
OUTPUT_DIR = SOURCE_DIR / "deliverables"
ASSET_DIR = SOURCE_DIR / "assets" / "technical-current"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
NAVY = "071724"
MAGENTA = "EC0A6B"
WHITE = "FFFFFF"
MID_GREY = "667085"
LIGHT_GREY = "F4F6F9"


@dataclass(frozen=True)
class Deliverable:
    source: Path
    output: Path
    title: str
    subtitle: str
    document_code: str
    preset: str
    cover_note: str
    toc: bool = False


DELIVERABLES = (
    Deliverable(
        source=SOURCE_DIR / "TECHNICAL_DOCUMENT.md",
        output=OUTPUT_DIR / "01_Dokumen_Teknis_Instalasi_dan_Penggunaan_HAWK-EYE.docx",
        title="HAWK-EYE — JudolGraph",
        subtitle="Dokumen Teknis: Panduan Instalasi dan Penggunaan",
        document_code="DT-01",
        preset="compact_reference_guide",
        cover_note="Struktur Gemastik a–e • maksimum 30 halaman • fixture aman",
        toc=True,
    ),
    Deliverable(
        source=SOURCE_DIR / "LIBRARIES_AND_LICENSES.md",
        output=OUTPUT_DIR / "05_Daftar_Komponen_dan_Lisensi_HAWK-EYE.docx",
        title="Daftar Komponen dan Lisensi",
        subtitle="HAWK-EYE — JudolGraph 1.0.0",
        document_code="DL-05",
        preset="compact_reference_guide",
        cover_note="Inventaris berbasis uv.lock dan pnpm-lock.yaml",
    ),
    Deliverable(
        source=SOURCE_DIR / "LICENSE_ADOPTION.md",
        output=OUTPUT_DIR / "07_Adopsi_Lisensi_HAWK-EYE.docx",
        title="Adopsi Lisensi",
        subtitle="MIT License untuk karya asli HAWK-EYE",
        document_code="AL-07",
        preset="standard_business_brief",
        cover_note="Tanggal berlaku 13 Agustus 2026",
    ),
)


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, keep: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if keep and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    elif not keep and node is not None:
        p_pr.remove(node)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph) -> None:
    paragraph.add_run("Halaman ")
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr, separate, text, end))


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Daftar isi diperbarui saat dokumen dibuka atau dirender."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, placeholder, end))


def set_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((run_color, underline))
    run.append(run_properties)
    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_markup(paragraph, text: str) -> None:
    """Render minimal Markdown emphasis/code and explicit web URLs."""
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|https?://[^\s)]+|(?:pypi|npm)\.org/[^\s|]+)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = rgb(DEEP_BLUE)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith(("http://", "https://")):
            add_hyperlink(paragraph, token, token)
        else:
            url = f"https://{token}"
            add_hyperlink(paragraph, token, url)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def set_run_font(run, font_name: str = "Calibri", size: float | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)


def configure_styles(document: Document, preset: str) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.10

    heading_values = {
        "Heading 1": (
            16,
            BLUE,
            18 if preset == "compact_reference_guide" else 16,
            10 if preset == "compact_reference_guide" else 8,
        ),
        "Heading 2": (
            13,
            BLUE,
            14 if preset == "compact_reference_guide" else 12,
            7 if preset == "compact_reference_guide" else 6,
        ),
        "Heading 3": (
            12,
            DEEP_BLUE,
            10 if preset == "compact_reference_guide" else 8,
            5 if preset == "compact_reference_guide" else 4,
        ),
    }
    for style_name, (size, color, before, after) in heading_values.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(
            0.375 if preset == "compact_reference_guide" else 0.5
        )
        style.paragraph_format.first_line_indent = Inches(
            -0.188 if preset == "compact_reference_guide" else -0.25
        )
        style.paragraph_format.space_after = Pt(4 if preset == "compact_reference_guide" else 8)
        style.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.167


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def add_cover(document: Document, spec: Deliverable) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.9)
    table.columns[1].width = Inches(1.9)
    left, right = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    set_cell_shading(left, NAVY)
    set_cell_shading(right, NAVY)
    set_cell_margins(left, top=180, start=240, bottom=180, end=80)
    set_cell_margins(right, top=100, start=80, bottom=100, end=160)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("HAWK-EYE")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = rgb(WHITE)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ALAT INVESTIGASI EKOSISTEM JUDI ONLINE")
    run.bold = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb("8BB7D0")

    avatar = ROOT / "apps" / "web" / "src" / "assets" / "hawkeye-avatar.png"
    image_run = right.paragraphs[0].add_run()
    image_run.add_picture(str(avatar), width=Inches(0.95))
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_alt_text(
        document.inline_shapes[-1],
        "Logo HAWK-EYE",
        "Ikon burung elang HAWK-EYE berwarna hitam dan magenta.",
    )

    stripe = document.add_table(rows=1, cols=1)
    stripe.alignment = WD_TABLE_ALIGNMENT.CENTER
    stripe.cell(0, 0).width = Inches(6.92)
    set_repeat_table_header(stripe.rows[0])
    set_cell_shading(stripe.cell(0, 0), MAGENTA)
    set_cell_margins(stripe.cell(0, 0), top=18, start=0, bottom=18, end=0)
    stripe.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)

    document.add_paragraph().paragraph_format.space_after = Pt(28)
    code = document.add_paragraph()
    code.paragraph_format.space_after = Pt(12)
    run = code.add_run(f"GEMASTIK 2026  /  {spec.document_code}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(MAGENTA)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(spec.title)
    run.bold = True
    run.font.size = Pt(30 if len(spec.title) < 28 else 27)
    run.font.color.rgb = rgb(NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(spec.subtitle)
    run.font.size = Pt(15)
    run.font.color.rgb = rgb(DEEP_BLUE)

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(42)
    run = note.add_run(spec.cover_note)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb(MAGENTA)

    metadata = document.add_table(rows=4, cols=2)
    set_repeat_table_header(metadata.rows[0])
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    metadata.autofit = False
    labels = ("Produk", "Versi", "Snapshot", "Klasifikasi")
    values = ("HAWK-EYE — JudolGraph", "1.0.0", "13 Agustus 2026", "Deliverable kompetisi")
    for row, label, value in zip(metadata.rows, labels, values, strict=True):
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(5.15)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label.upper())
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = rgb(DEEP_BLUE)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        r.font.size = Pt(9.5)
        r.font.color.rgb = rgb(NAVY)

    document.add_paragraph().paragraph_format.space_after = Pt(18)
    boundary = document.add_paragraph()
    boundary.paragraph_format.space_after = Pt(0)
    boundary.paragraph_format.left_indent = Inches(0.15)
    boundary.paragraph_format.right_indent = Inches(0.15)
    run = boundary.add_run(
        "Batas interpretasi — kandidat adalah lead, kemiripan bukan probabilitas "
        "kepemilikan, dan indikator adalah hitungan item bukti."
    )
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(MID_GREY)

    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run(f"HAWK-EYE  /  GEMASTIK 2026  /  {spec.document_code}")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(DEEP_BLUE)
    p_pr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), MAGENTA)
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run(f"{spec.document_code}  •  13 AGUSTUS 2026  •  ")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = rgb(MID_GREY)
    add_page_field(footer)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(MID_GREY)


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"-{3,}", line.strip()))


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    available = 6.92
    widths = [available / col_count] * col_count
    if col_count == 2:
        widths = [1.5, 5.42]
    elif col_count == 3:
        widths = [1.75, 1.2, 3.97]
    elif col_count == 4:
        widths = [1.45, 1.0, 2.7, 1.77]
    elif col_count == 5:
        widths = [1.28, 0.78, 2.1, 1.2, 1.56]
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        set_row_cant_split(row)
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GREY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if column_index < len(values):
                add_inline_markup(p, values[column_index])
            for run in p.runs:
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = rgb(DEEP_BLUE)
        if row_index == 0:
            set_repeat_table_header(row)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(document: Document, lines: Iterable[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    table.cell(0, 0).width = Inches(6.7)
    set_cell_shading(table.cell(0, 0), "EEF3F7")
    set_cell_margins(table.cell(0, 0), top=100, start=150, bottom=100, end=150)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 8.25)
    run.font.color.rgb = rgb(NAVY)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(document: Document, relative_path: str, alt_text: str) -> None:
    image_path = SOURCE_DIR / relative_path
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.7))
    set_alt_text(document.inline_shapes[-1], image_path.stem, alt_text)


def add_markdown_content(document: Document, source: Path, skip_headings: int, toc: bool) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    heading_count = 0
    index = 0
    if toc:
        heading = document.add_paragraph("Daftar isi", style="Heading 1")
        heading.paragraph_format.space_before = Pt(0)
        toc_p = document.add_paragraph()
        add_toc_field(toc_p)
        toc_p.add_run().add_break(WD_BREAK.PAGE)

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if is_separator(line):
            index += 1
            continue
        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(document, code_lines)
            continue
        image_match = re.fullmatch(r"!\[(.+)]\((.+)\)", line)
        if image_match:
            add_image(document, image_match.group(2), image_match.group(1))
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, rows)
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            heading_count += 1
            if heading_count <= skip_headings:
                index += 1
                continue
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline_markup(paragraph, text)
            index += 1
            continue
        bullet_match = re.match(r"^-\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.34)
            paragraph.paragraph_format.first_line_indent = Inches(-0.20)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.18
            paragraph.add_run("•  ")
            add_inline_markup(paragraph, bullet_match.group(1))
            index += 1
            continue
        number_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if number_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.40)
            paragraph.paragraph_format.first_line_indent = Inches(-0.28)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.18
            paragraph.add_run(f"{number_match.group(1)}.  ")
            add_inline_markup(paragraph, number_match.group(2))
            index += 1
            continue
        quote_match = re.match(r"^>\s?(.+)$", line)
        if quote_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            add_inline_markup(paragraph, quote_match.group(1))
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = rgb(DEEP_BLUE)
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", "- ", "```", "![", ">", "|"))
                or is_separator(candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_inline_markup(paragraph, " ".join(paragraph_lines))


def set_document_properties(document: Document, spec: Deliverable) -> None:
    properties = document.core_properties
    properties.title = spec.title
    properties.subject = spec.subtitle
    properties.author = "HAWK-EYE contributors"
    properties.keywords = "GEMASTIK 2026, HAWK-EYE, JudolGraph, evidence, OSINT"
    properties.comments = (
        f"Generated from {spec.source.relative_to(ROOT)} by a repository-tracked builder."
    )


def build(spec: Deliverable) -> None:
    document = Document()
    configure_page(document)
    configure_styles(document, spec.preset)
    set_document_properties(document, spec)
    add_cover(document, spec)
    skip_headings = 2 if spec.source.name == "TECHNICAL_DOCUMENT.md" else 1
    add_markdown_content(document, spec.source, skip_headings=skip_headings, toc=spec.toc)
    document.settings.update_fields_on_open = True
    spec.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(spec.output)
    print(f"Built {spec.output.relative_to(ROOT)}")


def main() -> int:
    for spec in DELIVERABLES:
        build(spec)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
